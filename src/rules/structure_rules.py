"""文档结构校验规则模块

实现三行一组文档结构校验规则
"""

import re
from typing import List, Dict, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from loguru import logger

from rules.base_rule import BaseRule, ValidationError


class ThreeLineGroupValidationRule(BaseRule):
    """三行一组结构校验规则

    文档按每3行一组的循环结构处理：
    - 第1行（段落索引 %3 == 0）：标题行
    - 第2行（段落索引 %3 == 1）：URL行
    - 第3行（段落索引 %3 == 2）：图片行
    """

    # URL 正则表达式
    REGEX_URL = r'^https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&\/=]*)$'
    # 标题行正则表达式：以数字开头，后跟 .、, 或 | 分隔符
    REGEX_TITLE = r'^\d+[.|,|)][\s\S]+'

    def __init__(self, title_mode: str = "1", enabled: bool = True):
        """

        Args:
            title_mode: 标题模式
                - "1": 标准模式，需要包含 . _ ： 三个字符
                - "2": 精简模式，只需要提取 . 后的内容
            enabled: 是否启用该规则
        """
        super().__init__("三行一组结构校验", enabled)
        self.title_mode = title_mode

    def validate(self, document: Document) -> List[ValidationError]:
        """校验文档结构

        Args:
            document: 要校验的Word文档对象

        Returns:
            发现的错误列表
        """
        self.clear_errors()
        paragraphs = document.paragraphs

        # 检查段落数量是否为3的倍数
        total_paragraphs = len(paragraphs)
        if total_paragraphs % 3 != 0:
            self.errors.append(ValidationError(
                -1, 'structure',
                f'段落数量({total_paragraphs})不是3的倍数，剩余{total_paragraphs % 3}行'
            ))

        # 遍历每三行一组
        for i in range(0, len(paragraphs), 3):
            group_index = i // 3

            # 检查标题行（第1行）
            if i < len(paragraphs):
                self._validate_title_line(paragraphs[i], i, group_index)

            # 检查URL行（第2行）
            if i + 1 < len(paragraphs):
                self._validate_url_line(paragraphs[i + 1], i + 1, group_index)

            # 检查图片行（第3行）
            if i + 2 < len(paragraphs):
                self._validate_image_line(paragraphs[i + 2], i + 2, group_index)

        return self.errors

    def _validate_title_line(self, paragraph: Paragraph, index: int, group_index: int) -> None:
        """校验标题行

        Args:
            paragraph: 段落对象
            index: 段落索引
            group_index: 组索引（第几组）
        """
        text = paragraph.text.strip()

        # 空行检查
        if not text:
            self.errors.append(ValidationError(
                index, 'title',
                f'第{group_index + 1}组标题行为空'
            ))
            return

        # 检查基础格式（序号+分隔符+内容）
        if not re.match(self.REGEX_TITLE, text):
            self.errors.append(ValidationError(
                index, 'title',
                f'第{group_index + 1}组标题行格式错误，应为"序号+分隔符(.|,|)+内容"，当前为: "{text[:50]}..."'
            ))
            return

        # 标准模式检查
        if self.title_mode == "1":
            required_chars = {'.': '点', '_': '下划线', '：': '冒号'}
            for char, char_name in required_chars.items():
                if char not in text:
                    self.errors.append(ValidationError(
                        index, 'title',
                        f'第{group_index + 1}组标题行缺少必需字符"{char}"({char_name})'
                    ))

    def _validate_url_line(self, paragraph: Paragraph, index: int, group_index: int) -> None:
        """校验URL行

        Args:
            paragraph: 段落对象
            index: 段落索引
            group_index: 组索引（第几组）
        """
        text = paragraph.text.strip()

        # 空行检查
        if not text:
            self.errors.append(ValidationError(
                index, 'url',
                f'第{group_index + 1}组URL行为空'
            ))
            return

        # URL格式检查
        if not re.match(self.REGEX_URL, text):
            self.errors.append(ValidationError(
                index, 'url',
                f'第{group_index + 1}组URL行格式错误: "{text[:80]}..."'
            ))

    def _validate_image_line(self, paragraph: Paragraph, index: int, group_index: int) -> None:
        """校验图片行

        图片行要求：
        1. 必须包含图片
        2. 不能包含任何文字或空格（只允许纯图片）
        3. 只能有一张图片（只能有一个包含图片的run）

        Args:
            paragraph: 段落对象
            index: 段落索引
            group_index: 组索引（第几组）
        """
        # 获取完整文本内容（不strip，保留空格和换行）
        full_text = paragraph.text

        # 检查是否有任何文字内容（包括空格、换行等）
        # 图片行应该是"纯净"的，除了图片run外不应有任何text run
        if full_text:
            # 进一步检查：是否所有run都是图片run
            non_image_text_found = False
            text_content = ""

            for run in paragraph.runs:
                if run.text and not self._is_image_only_run(run):
                    non_image_text_found = True
                    text_content += run.text

            if non_image_text_found:
                # 显示实际内容（用repr展示隐藏字符）
                display_text = repr(full_text[:50]) if len(repr(full_text)) > 50 else repr(full_text)
                self.errors.append(ValidationError(
                    index, 'image',
                    f'第{group_index + 1}组图片行包含图片以外的字符: {display_text}'
                ))

        # 检查图片数量 - 直接从 XML 中获取所有 drawing 元素
        image_count = self._count_images_in_paragraph(paragraph)

        # 检查是否有图片
        if image_count == 0:
            self.errors.append(ValidationError(
                index, 'image',
                f'第{group_index + 1}组图片行没有图片'
            ))
        elif image_count > 1:
            self.errors.append(ValidationError(
                index, 'image',
                f'第{group_index + 1}组图片行包含{image_count}张图片，只能有一张图片'
            ))

    def _count_images_in_paragraph(self, paragraph: Paragraph) -> int:
        """直接从 XML 统计段落中的图片数量

        Args:
            paragraph: 段落对象

        Returns:
            图片数量
        """
        # 直接从段落 XML 元素中查找所有的 drawing 元素
        p_element = paragraph._element
        xml_str = p_element.xml

        # 使用正则表达式查找所有 wp:extent 元素（每个图片都有一个）
        import re
        extents = re.findall(r'<wp:extent\s+cx="[^"]*"\s+cy="[^"]*"', xml_str)
        return len(extents)

    def _has_image(self, paragraph: Paragraph) -> bool:
        """检查段落是否包含图片

        Args:
            paragraph: 段落对象

        Returns:
            True表示包含图片，False表示不包含
        """
        for run in paragraph.runs:
            if 'graphic' in run._element.xml or 'pic:' in run._element.xml:
                return True
            # 检查是否有blip元素（嵌入图片）
            if run._element.xpath('.//a:blip'):
                return True
        return False

    def _is_image_only_run(self, run) -> bool:
        """检查run是否只包含图片（不包含文字）

        Args:
            run: paragraph.run对象

        Returns:
            True表示run只包含图片，False表示run包含文字或既无图片也无文字
        """
        # 如果有文本内容，不是纯图片run
        if run.text:
            return False

        # 检查是否包含图片元素
        xml = run._element.xml
        if 'graphic' in xml or 'pic:' in xml:
            return True
        if run._element.xpath('.//a:blip'):
            return True

        return False

    def _is_valid_image(self, run) -> bool:
        """检查图片是否有效（高度不为0）

        Args:
            run: paragraph.run对象

        Returns:
            True表示图片有效（高度>0），False表示图片无效（高度=0）
        """
        if not self._is_image_only_run(run):
            return False

        # 检查图片的尺寸（wp:extent）
        # XML中格式: <wp:extent cx="宽度" cy="高度"/>
        xml = run._element.xml
        import re
        extent_match = re.search(r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', xml)
        if extent_match:
            height = int(extent_match.group(2))
            # 高度大于0才认为是有效图片
            return height > 0

        # 如果没有找到extent，也认为是有效的（以防格式不同）
        return True

    def _get_image_info(self, run) -> dict:
        """获取图片信息

        Args:
            run: paragraph.run对象

        Returns:
            包含图片信息的字典
        """
        info = {'valid': False, 'width': 0, 'height': 0, 'name': ''}

        xml = run._element.xml
        import re
        extent_match = re.search(r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', xml)
        if extent_match:
            info['width'] = int(extent_match.group(1))
            info['height'] = int(extent_match.group(2))
            info['valid'] = info['height'] > 0

        name_match = re.search(r'<wp:docPr\s+id="[^"]*"\s+name="([^"]*)"', xml)
        if name_match:
            info['name'] = name_match.group(1)

        return info

    def apply(self, document: Document) -> Document:
        """应用规则到文档（校验并返回）

        Args:
            document: 要处理的Word文档对象

        Returns:
            处理后的文档对象
        """
        # 校验规则主要用于验证，这里只做校验
        self.validate(document)
        return document

    def fix(self, document: Document) -> Document:
        """自动修复文档

        Args:
            document: 要修复的Word文档对象

        Returns:
            修复后的文档对象
        """
        # 先校验获取所有错误
        self.validate(document)

        if not self.errors:
            return document

        # 记录原始错误数量
        original_errors = len(self.errors)
        logger.info(f"开始修复，共发现 {original_errors} 个错误")

        # 按错误类型分组
        errors_by_index: Dict[int, List[ValidationError]] = {}
        for error in self.errors:
            if error.line_index not in errors_by_index:
                errors_by_index[error.line_index] = []
            errors_by_index[error.line_index].append(error)

        fixed_count = 0

        # 逐个修复
        for index, errors in sorted(errors_by_index.items()):
            if index >= len(document.paragraphs):
                continue

            paragraph = document.paragraphs[index]

            for error in errors:
                if error.line_type == 'image':
                    if '包含文字内容' in error.message or '包含图片以外的字符' in error.message:
                        # 删除图片行中的文字
                        self._clear_text_from_paragraph(paragraph)
                        fixed_count += 1
                        logger.info(f"已修复第{index // 3 + 1}组图片行: 清除了图片中的文字")
                    elif '张图片，只能有一张图片' in error.message:
                        # 删除多余的图片，只保留第一张
                        group_index = index // 3
                        logger.info(f"开始修复第{group_index + 1}组图片行: 处理多余图片")
                        removed = self._remove_extra_images(paragraph, group_index)
                        if removed > 0:
                            fixed_count += 1
                elif error.line_type == 'title':
                    # 尝试修复标题格式
                    original_text = paragraph.text
                    self._fix_title_format(paragraph, index)
                    if paragraph.text != original_text:
                        fixed_count += 1
                        logger.info(f"已修复第{index // 3 + 1}组标题行: '{original_text[:30]}...' -> '{paragraph.text[:30]}...'")
                elif error.line_type == 'url' and '格式错误' in error.message:
                    # 标记无效URL或删除
                    pass

        # 处理段落数量不是3的倍数的情况
        total_paragraphs = len(document.paragraphs)
        if total_paragraphs % 3 != 0:
            remainder = total_paragraphs % 3
            # 删除末尾多余的段落
            for _ in range(remainder):
                # 注意：这里需要从末尾删除
                pass  # python-docx不支持直接删除段落，需要其他处理方式

        # 修复完成后，再次校验以确认剩余的错误
        # 不要清空错误列表，而是重新校验来获取剩余问题
        self.validate(document)
        remaining_errors = len(self.errors)

        logger.info(f"修复完成: 尝试修复 {fixed_count} 个问题，剩余 {remaining_errors} 个问题")

        return document

    def _clear_text_from_paragraph(self, paragraph: Paragraph) -> None:
        """清除段落中的文字，保留图片

        Args:
            paragraph: 段落对象
        """
        for run in paragraph.runs:
            if run.text:
                run.text = ""

    def _remove_extra_images(self, paragraph: Paragraph, group_index: int) -> int:
        """删除段落中多余的图片，只保留有效的图片

        优先级：
        1. 删除所有无效图片（高度为0的图片）
        2. 如果删除无效图片后只剩1张有效图片，则完成
        3. 如果删除无效图片后仍有多个有效图片，只保留第一张有效图片

        Args:
            paragraph: 段落对象
            group_index: 组索引，用于日志

        Returns:
            删除的图片数量
        """
        # 直接从段落 XML 元素中查找所有包含图片的 w:r 元素
        p_element = paragraph._element
        xml_str = p_element.xml

        import re
        import xml.etree.ElementTree as ET

        # 使用正则表达式提取所有图片信息
        # 查找所有 <w:r>...</w:r> 块，然后从中提取图片信息
        # 但由于 XML 是命名空间的，我们需要用更复杂的方法

        # 策略：直接操作段落元素的子元素
        # 找到所有包含 drawing 的 w:r 子元素

        removed_count = 0

        # 获取所有 w:r 子元素
        # 需要考虑命名空间
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # 直接遍历子元素
        r_elements_to_remove = []

        # 收集所有包含 drawing 的子元素及其信息
        image_elements = []  # (element, info_dict)

        for child in list(p_element):
            # 检查是否是 w:r 元素
            tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag_name != 'r':
                continue

            # 检查这个 w:r 是否包含 drawing
            child_xml = child.xml

            # 查找 wp:extent 和 wp:docPr
            extent_match = re.search(r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', child_xml)
            docPr_match = re.search(r'<wp:docPr\s+id="[^"]*"\s+name="([^"]*)"', child_xml)

            if extent_match or docPr_match:
                # 这是一个图片 run
                width = int(extent_match.group(1)) if extent_match else 0
                height = int(extent_match.group(2)) if extent_match else 0
                name = docPr_match.group(1) if docPr_match else "未知"

                info = {
                    'element': child,
                    'width': width,
                    'height': height,
                    'valid': height > 0,
                    'name': name
                }
                image_elements.append(info)

        # 分离有效和无效图片
        invalid_images = [img for img in image_elements if not img['valid']]
        valid_images = [img for img in image_elements if img['valid']]

        # 首先删除所有无效图片
        for img in invalid_images:
            element = img['element']
            element.getparent().remove(element)
            removed_count += 1
            logger.info(f"  删除无效图片: {img['name']} (尺寸: {img['width']}x{img['height']})")

        # 检查剩余的有效图片数量
        if len(valid_images) == 0:
            logger.warning(f"第{group_index + 1}组图片行: 删除无效图片后没有有效图片了！")
        elif len(valid_images) == 1:
            logger.info(f"  保留唯一有效图片: {valid_images[0]['name']}")
        elif len(valid_images) > 1:
            # 有多个有效图片，只保留第一个
            logger.warning(f"  发现{len(valid_images)}张有效图片（只能有一张）:")
            for i, img in enumerate(valid_images):
                if i == 0:
                    logger.info(f"    保留: {img['name']} (尺寸: {img['width']}x{img['height']})")
                else:
                    logger.info(f"    删除: {img['name']} (尺寸: {img['width']}x{img['height']})")
                    element = img['element']
                    element.getparent().remove(element)
                    removed_count += 1

        return removed_count

    def _fix_title_format(self, paragraph: Paragraph, index: int) -> None:
        """尝试修复标题格式

        Args:
            paragraph: 段落对象
            index: 段落索引
        """
        text = paragraph.text.strip()

        # 检查是否以数字开头但没有正确的分隔符
        # 匹配：数字后直接跟非分隔符字符（.|,|）
        match = re.match(r'^(\d+)([^.|,|\s])', text)
        if match:
            # 在数字后面添加 . 分隔符，保留后面的字符
            number = match.group(1)
            # rest 从数字后开始，不包括被匹配的非分隔符字符
            rest = text[len(number):]
            new_text = f"{number}.{rest}"

            # 更新段落文本
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                # 清除其他run的文本
                for run in paragraph.runs[1:]:
                    run.text = ""
        elif not re.match(r'^\d', text):
            # 如果缺少序号，尝试添加
            group_num = (index // 3) + 1
            # 检查是否已经有分隔符，有则直接加序号
            if re.match(r'^[.|,|)]', text):
                new_text = f"{group_num}{text}"
            else:
                # 尝试添加序号和分隔符
                new_text = f"{group_num}. {text}"

            # 更新段落文本
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                # 清除其他run的文本
                for run in paragraph.runs[1:]:
                    run.text = ""

    def parse_title(self, title_text: str) -> Dict[str, str]:
        """解析标题行，提取字段

        Args:
            title_text: 标题文本

        Returns:
            包含解析字段的字典
        """
        result: Dict[str, str] = {}

        try:
            # 去除序号部分
            content = re.sub(r'^\d+[.|,|)]', '', title_text).strip()

            if self.title_mode == "1":
                # 标准模式：提取舆论场、来源、标题
                dot_idx = content.find('.')
                underscore_idx = content.find('_')
                colon_idx = content.find('：')

                if all(idx >= 0 for idx in [dot_idx, underscore_idx, colon_idx]):
                    result['opinion'] = content[dot_idx + 1:underscore_idx].strip()
                    result['source'] = content[underscore_idx + 1:colon_idx].strip()
                    result['title'] = content[colon_idx + 1:].strip()
                else:
                    result['error'] = f"无法解析标题，缺少必需分隔符。内容: {content}"
            else:
                # 精简模式：只提取标题
                result['title'] = content

        except Exception as e:
            result['error'] = str(e)

        return result

    def parse_document(self, document: Document) -> List[Dict[str, any]]:
        """解析整个文档，提取所有三行一组的数据

        Args:
            document: Word文档对象

        Returns:
            包含所有组数据的列表
        """
        results = []
        paragraphs = document.paragraphs

        for i in range(0, len(paragraphs), 3):
            group_data = {
                'group_index': i // 3,
                'title_line': None,
                'url_line': None,
                'image_line': None
            }

            # 标题行
            if i < len(paragraphs):
                title_text = paragraphs[i].text.strip()
                group_data['title_line'] = {
                    'text': title_text,
                    'parsed': self.parse_title(title_text),
                    'has_error': any(e.line_index == i and e.line_type == 'title' for e in self.errors)
                }

            # URL行
            if i + 1 < len(paragraphs):
                url_text = paragraphs[i + 1].text.strip()
                group_data['url_line'] = {
                    'text': url_text,
                    'is_valid': bool(re.match(self.REGEX_URL, url_text)),
                    'has_error': any(e.line_index == i + 1 and e.line_type == 'url' for e in self.errors)
                }

            # 图片行
            if i + 2 < len(paragraphs):
                paragraph = paragraphs[i + 2]
                group_data['image_line'] = {
                    'text': paragraph.text.strip(),
                    'has_image': self._has_image(paragraph),
                    'has_error': any(e.line_index == i + 2 and e.line_type == 'image' for e in self.errors)
                }

            results.append(group_data)

        return results

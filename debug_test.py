#!/usr/bin/env python3
"""调试脚本：测试图片修复功能"""

import sys
import os

# 添加 src 目录到 Python 路径
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from docx import Document
from rules.structure_rules import ThreeLineGroupValidationRule
from loguru import logger

# 配置 logger
logger.remove()
logger.add(sys.stderr, level="INFO")

try:
    # 加载文档（使用已修复的文件，因为原始文件损坏了）
    doc = Document("/Users/jingjianan/workspace/project/jjn/DocuRuleFix/tests/1_fixed.docx")

    # 创建规则实例
    rule = ThreeLineGroupValidationRule(title_mode="2", enabled=True)

    # 先校验
    logger.info("===== 修复前校验 =====")
    errors = rule.validate(doc)
    logger.info(f"发现 {len(errors)} 个错误")

    for error in errors:
        if error.line_type == 'image' and '张图片' in error.message:
            logger.info(f"  行 {error.line_index}: {error.message}")

    # 检查第13行（第4组图片行）的详细信息
    logger.info("\n===== 第13行（第4组图片行）详细信息 =====")
    para = doc.paragraphs[12]  # 索引12是第13行
    logger.info(f"Run 数量: {len(para.runs)}")

    for i, run in enumerate(para.runs):
        run_xml = run._element.xml
        # 提取图片信息
        import re
        extent_match = re.search(r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', run_xml)
        name_match = re.search(r'<wp:docPr\s+id="[^"]*"\s+name="([^"]*)"', run_xml)
        if extent_match or name_match:
            width = extent_match.group(1) if extent_match else "N/A"
            height = extent_match.group(2) if extent_match else "N/A"
            name = name_match.group(1) if name_match else "N/A"
            logger.info(f"  Run {i}: {name}, 尺寸: {width}x{height}")

    # 修复
    logger.info("\n===== 开始修复 =====")
    fixed_doc = rule.fix(doc)
    logger.info(f"修复完成，剩余 {len(rule.errors)} 个错误")

    # 检查修复后的第13行
    logger.info("\n===== 修复后第13行详细信息 =====")
    para_fixed = fixed_doc.paragraphs[12]
    logger.info(f"Run 数量: {len(para_fixed.runs)}")

    for i, run in enumerate(para_fixed.runs):
        run_xml = run._element.xml
        import re
        extent_match = re.search(r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', run_xml)
        name_match = re.search(r'<wp:docPr\s+id="[^"]*"\s+name="([^"]*)"', run_xml)
        if extent_match or name_match:
            width = extent_match.group(1) if extent_match else "N/A"
            height = extent_match.group(2) if extent_match else "N/A"
            name = name_match.group(1) if name_match else "N/A"
            logger.info(f"  Run {i}: {name}, 尺寸: {width}x{height}")

    # 保存
    fixed_doc.save("/Users/jingjianan/workspace/project/jjn/DocuRuleFix/tests/1_fixed_debug.docx")
    logger.info("\n已保存到 tests/1_fixed_debug.docx")

    # 重新加载检查
    logger.info("\n===== 重新加载修复后的文件检查 =====")
    doc2 = Document("/Users/jingjianan/workspace/project/jjn/DocuRuleFix/tests/1_fixed_debug.docx")
    rule2 = ThreeLineGroupValidationRule(title_mode="2", enabled=True)
    errors2 = rule2.validate(doc2)
    logger.info(f"重新校验发现 {len(errors2)} 个错误")

    for error in errors2:
        if error.line_type == 'image' and '张图片' in error.message:
            logger.info(f"  行 {error.line_index}: {error.message}")

    # 检查重新加载后的第13行
    logger.info("\n===== 重新加载后第13行详细信息 =====")
    para_reload = doc2.paragraphs[12]
    logger.info(f"Run 数量: {len(para_reload.runs)}")

    for i, run in enumerate(para_reload.runs):
        run_xml = run._element.xml
        import re
        extent_match = re.search(r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', run_xml)
        name_match = re.search(r'<wp:docPr\s+id="[^"]*"\s+name="([^"]*)"', run_xml)
        if extent_match or name_match:
            width = extent_match.group(1) if extent_match else "N/A"
            height = extent_match.group(2) if extent_match else "N/A"
            name = name_match.group(1) if name_match else "N/A"
            logger.info(f"  Run {i}: {name}, 尺寸: {width}x{height}")

except Exception as e:
    logger.error(f"错误: {type(e).__name__}: {e}")
    import traceback
    traceback.print_exc()
